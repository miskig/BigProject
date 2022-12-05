import sys
import subprocess
from docx import Document

if len(sys.argv) != 2:
    print("Please provide a file as an argument")

filename = sys.argv[1]

# check if file is a doc or docx file
if not filename.endswith(".doc") and not filename.endswith(".docx"):
    print("Unsupported file type. Only doc and docx files are supported.")


# extract text from doc or docx file
if filename.endswith(".doc"):
    doc = Document(filename)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
else: # filename.endswith(".docx")
    doc = docx.Document(filename)
    text = doc.core_properties.content_status

# call another script with filename and text as arguments
subprocess.call(["other_script.py", filename, text])
